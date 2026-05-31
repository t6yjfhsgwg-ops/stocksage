"""Convert docs/DESIGN.md to Word (.docx)."""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_LINE_SPACING
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_LINE_SPACING


def set_normal_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(6)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(line.rstrip())
        p.style = "No Spacing"
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def parse_table_row(line):
    line = line.strip()
    if not line.startswith("|"):
        return None
    cells = [c.strip() for c in line.strip("|").split("|")]
    return cells


def is_separator_row(cells):
    if not cells:
        return False
    return all(re.match(r"^:?-+:?$", c.replace(" ", "")) or c == "" for c in cells)


def md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    set_normal_style(doc)

    # Title
    title = doc.add_heading("StockSage — Design Document", 0)
    title.runs[0].font.name = "Calibri Light"

    meta_done = False
    i = 0
    in_code = False
    code_buf = []
    code_lang = ""
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if len(table_buf) < 1:
            table_buf = []
            return
        rows = []
        for row_line in table_buf:
            cells = parse_table_row(row_line)
            if cells and not is_separator_row(cells):
                rows.append(cells)
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    t.rows[ri].cells[ci].text = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell)
            doc.add_paragraph()
        table_buf = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                if code_lang == "mermaid":
                    p = doc.add_paragraph("Architecture diagram (Mermaid — render in GitHub or mermaid.live):")
                    p.runs[0].bold = True
                elif code_lang:
                    p = doc.add_paragraph(f"Code ({code_lang}):")
                    p.runs[0].bold = True
                add_code_block(doc, code_buf)
                doc.add_paragraph()
                in_code = False
                code_buf = []
                code_lang = ""
            else:
                code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            code_lang = line.strip()[3:].strip() or "text"
            i += 1
            continue

        if line.strip().startswith("|"):
            flush_table()
            table_buf.append(line)
            i += 1
            continue
        else:
            flush_table()

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), 1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), 2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), 3)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), 4)
            i += 1
            continue

        if line.strip().startswith("- "):
            item = line.strip()[2:]
            item = re.sub(r"\*\*([^*]+)\*\*", r"\1", item)
            item = item.replace("&lt;", "<").replace("&gt;", ">")
            doc.add_paragraph(item, style="List Bullet")
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.+)$", line.strip())
        if m:
            item = re.sub(r"\*\*([^*]+)\*\*", r"\1", m.group(2))
            doc.add_paragraph(item, style="List Number")
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # Skip duplicate title from md
        if line.startswith("# StockSage") and meta_done:
            i += 1
            continue

        para = line.strip()
        para = re.sub(r"\*\*([^*]+)\*\*", r"\1", para)
        para = re.sub(r"`([^`]+)`", r"\1", para)
        para = para.replace("&lt;", "<").replace("&gt;", ">")
        p = doc.add_paragraph(para)
        if not meta_done and para.startswith("StockSage"):
            meta_done = True
        i += 1

    flush_table()
    if in_code and code_buf:
        add_code_block(doc, code_buf)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "docs" / "DESIGN.md"
    out = root / "docs" / "StockSage-Design-Document.docx"
    if len(sys.argv) > 1:
        out = Path(sys.argv[1])
    md_to_docx(md, out)
